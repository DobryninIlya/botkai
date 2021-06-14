from .. import classes
from .. import keyboards


import sqlite3
import datetime
import json
import random 
import requests
import traceback
import os, importlib
import sys
import apiai
from pprint import pprint
from numba import jit

from bs4 import BeautifulSoup
import docx
from docx.shared import Inches, Cm
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE,WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


cursor = classes.cursor
cursorR = classes.cursorR
conn = classes.conn
connection = classes.connection

try:
    cursorR.execute("""CREATE TABLE storage (id INT NOT NULL PRIMARY KEY, media_id INT NOT NULL); """)
    cursorR.execute("""CREATE TABLE answers (id INT NOT NULL PRIMARY KEY, userId INT NOT NULL); """)
    cursorR.execute("""CREATE TABLE prepod_users (id INT NOT NULL PRIMARY KEY, groupreal INT NOT NULL, groupid INT NOT NULL); """)
    cursorR.execute("""CREATE TABLE Status (ID_VK INT NOT NULL PRIMARY KEY, Status INT NULL); """)
    conn.commit()
except:
    pass
today = datetime.date.today()
message_params = {}

vk = classes.vk

UserParams = classes.UserParams
command_list = classes.command_list

MessageSettings=classes.MessageSettings

def load_modules():
    try:
        files = os.listdir("/app/botkai/commands")
    except:
        files = os.listdir("botkai/commands")
    modules = filter(lambda x: x.endswith('.py'), files)
    for m in modules:
        importlib.import_module("botkai.commands." + m[0:-3])


load_modules()

def DeleteOldTask():
    cursor.execute("SELECT COUNT(*) FROM Task WHERE Datee < '" + str(datetime.date(today.year, today.month, today.day) - datetime.timedelta(days=1)) + "'")
    count = cursor.fetchone()[0]
    cursor.execute("SELECT * FROM Task WHERE Datee < '" + str(datetime.date(today.year, today.month, today.day) - datetime.timedelta(days=1)) + "'")

    result = cursor.fetchall()
    #print(result)
    for part in result:
        try:
            vk.method("messages.send", {"peer_id": (int)(str(part[2])), "message": "Ваше задание " + (str(part[4]))[:10] + "... было автоматически удалено.", "keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
        except Exception as E:
            pass
    cursor.execute("DELETE FROM Task WHERE Datee < '" + str(datetime.date(today.year, today.month, today.day)  - datetime.timedelta(days=1)) + "'")
    cursor.execute('DELETE FROM "Adv" WHERE date < ' + "'" + str(datetime.date(today.year, today.month, today.day)  - datetime.timedelta(days=1)) + "'")

    cursor.execute("DELETE FROM Users WHERE Groupp = 0 and role = 1")


    connection.commit()
    if count == 0:
        return
    vk.method("messages.send", {"peer_id": 159773942, "message": "Удалено заданий: " + str(count) , "keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
    
DeleteOldTask()


def textMessage():
    try:
        request = apiai.ApiAI('07ed539ee928497ebb68729e87c09ede').text_request() # Токен API к Dialogflow
        request.lang = 'ru' # На каком языке будет послан запрос
        request.session_id = 'kaibot' # ID Сессии диалога (нужно, чтобы потом учить бота)
        request.query = MessageSettings.getText() # Посылаем запрос к ИИ с сообщением от юзера
        responseJson = json.loads(request.getresponse().read().decode('utf-8'))
        response = responseJson['result']['fulfillment']['speech'] # Разбираем JSON и вытаскиваем ответ
        # Если есть ответ от бота - присылаем юзеру, если нет - бот его не понял
        if response:
            mesg = response
            vk.method("messages.send",
                    {"peer_id": MessageSettings.getId(), "message": mesg,"keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
        else:
            mesg = "Я не совсем понял тебя."
            if not MessageSettings.keyboard:
                mesg+="\nУ тебя неофициальный клиент, введи 'Команды', чтобы получить список основных команд"
            vk.method("messages.send",
                    {"peer_id": MessageSettings.getId(), "message": mesg,"keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
    except Exception:
        vk.method("messages.send",
                {"peer_id": MessageSettings.getId(), "message": "Я не понял тебя","keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})


def message_new(request):

    try:
        global message_params
        message_params = json.loads(request.body)
        MessageSettings.update(message_params)
        if MessageSettings.secret_key != os.getenv("SECRET_KEY"):
            vk.method("messages.send",
                      {"peer_id": 159773942, "message": "Попытка взлома. Секретный ключ не сошелся\n{}".format(MessageSettings.secret_key), "random_id": random.randint(1, 2147483647)})
            print("СЕКРЕТНЫЙ КЛЮЧ НЕ СОШЕЛСЯ")
            return "У тебя почти получилось :)"
        if IsRegistred():

            if MessageSettings.peer_id > 2000000000:
                return "ok"

            UserParams.update(int(MessageSettings.id))
            UserParams.Status = StatusR(MessageSettings.getId())
            stat = CheckStatus()
            if stat == "ok":
                return "ok"
            
            cursorR.execute("SELECT * FROM Status")

            button = ""
            try:
                # print("msg payload", MessageSettings.payload)
                payload = MessageSettings.payload
                button = payload["button"]
                MessageSettings.button = button
            except Exception as E:
                pass

            if button != "":
                for c in command_list:
                    crole = c.role
                    if button == c.payload and c.admlevel<=UserParams.getAdminLevel():
                        if UserParams.role in crole:
                            c.process()
                        return "ok"
                return "ok"
            else:
                distance = len(MessageSettings.getText())
                command = None
                key = ''
                for c in command_list:
                    
                    if UserParams.role in c.role:
                        
                        for k in c.keys:
                            d = damerau_levenshtein_distance(( MessageSettings.getText()).lower(), k)
                            if d < distance:
                                distance = d
                                command = c
                                key = k
                                MessageSettings.command_key = k
                                if distance == 0 and c.admlevel<=UserParams.getAdminLevel() and (UserParams.role in c.role):
                                    c.process()
                                    return "ok"
                if distance < len(MessageSettings.getText())*0.4 and command.admlevel<=UserParams.getAdminLevel()  and (UserParams.role in command.role):
                    
                    mesg = 'Я понял ваш запрос как "%s"' % key
                    MessageSettings.command_key = key
                    vk.method("messages.send",
                            {"peer_id": int(MessageSettings.id), "message": mesg, "random_id": random.randint(1, 2147483647)})
                    command.process()
                    return "ok"
            textMessage()

    except SystemExit:
        quit()
        sys.exit(1)
        os.abort()
    except:  
        print('Ошибка:\n', traceback.format_exc())

    return "ok"





def IsRegistred():
    try:
        body = MessageSettings.getText()
        id = int(MessageSettings.id)
        payload = ""
        try:
            if MessageSettings.payload:
                if MessageSettings.payload["button"]:
                    payload = MessageSettings.payload["button"]
        except:
            pass
        if payload == "undo_regestration":
            sql = "UPDATE Status SET Status = 3 WHERE id_vk = {}".format(id)
            cursorR.execute(sql)
            conn.commit()
            vk.method("messages.send", {"peer_id": id, "message": "Мне нужно понимать кто ты. Выбери соответствующую кнопку в меню", "keyboard" : keyboards.get_undo,
                                    "random_id": random.randint(1, 2147483647)})

        if InBase(id):
            #print("Зарегистрироан")
            return True
        else:

            #print("Не зарегистрирован")
            if InBaseR(id):
                vk.method("messages.send",
                            {"peer_id": id, "message": "test" , "sticker_id" : 6864 , "random_id": random.randint(1, 2147483647)})
                vk.method("messages.send", {"peer_id": id, "message": "Похоже, что ты не зарегистриован. Для работы бота необходима регистрация.\nПо любым вопросам введите Справка в чат, чтобы продолжить пользоваться ботом - выполняйте инструкции.\n Мне нужно понимать кто ты. Выбери соответствующую кнопку в меню", "keyboard" : keyboards.roleMenu,
                                    "random_id": random.randint(1, 2147483647)})
                sql = "INSERT INTO Status VALUES (" + str(id) + ", 3);"
                cursorR.execute(sql)
                conn.commit()
                return False
            elif StatusR(id) == 3:
                
                today = datetime.date.today()
                role = 0
                if body == "Преподаватель":
                    role = 2
                elif body == "Студент":
                    role = 1
                elif body == "Родитель":
                    role = 3
                elif body == "Абитуриент (поступающий)":
                    role = 4
                elif body == "Справка":
                    msg = """На данном этапе необходимо указать свою роль. 
                    Если вы студент, вам будут доступно свое расписание, список преподавателей. Вас могут видеть одногруппники в списке группы.
                    Если вы родитель, то вам будет доступно расписание и список преподавателей. Ваш аккаунт будет скрыт от вашего ребенка. Вам также будет доступен список других родителей данной группы.
                    Если вы преподаватель, вам будет доступно ваше расписание и рассылка объявлений группе.
                    
                    """
                    vk.method("messages.send", {"peer_id": id, "message": msg, "keyboard" : keyboards.roleMenu, 
                                    "random_id": random.randint(1, 2147483647)})
                    return False
                else:
                    vk.method("messages.send", {"peer_id": id, "message": "Выберите кнопку меню.", "keyboard" : keyboards.roleMenu, 
                                    "random_id": random.randint(1, 2147483647)})
                    return False
                    
                try:
                    sql = "INSERT INTO Users (id_vk, name, groupp, distribution, admLevel, groupreal, \"dateChange\", balance, distr, warn, expiration, banhistory, ischeked, role, login, potok_lecture, has_own_shed, affiliate)" \
                          " VALUES (" + str(id) + ", '" + "', " + "0 " + ", 1, 1, 0, '" + str(datetime.date(today.year, today.month, today.day)) +"'"\
                          ",0 , 0, 0, '2020-01-01', 0, 0," + str(role) + ", null, true, false, false);"
                    cursor.execute(sql)
                    connection.commit()
                except Exception as E:
                    print('Ошибка commit:\n', traceback.format_exc())
                if role == 1 or role == 3:
                    
                    sql = "UPDATE Status SET Status = 1 WHERE ID_VK = " + str(id) + ";"
                    cursorR.execute(sql)
                    conn.commit()
                    vk.method("messages.send", {"peer_id": id, "message": "Введите свое имя в чат", "keyboard" : keyboards.get_undo,
                                    "random_id": random.randint(1, 2147483647)})
                elif role == 2:
                    sql = "UPDATE Status SET Status = 4 WHERE ID_VK = " + str(id) + ";"
                    cursorR.execute(sql)
                    conn.commit()
                    vk.method("messages.send", {"peer_id": id, "message": "Введите свой логин без лишних символов.", "keyboard" : keyboards.get_undo,
                                    "random_id": random.randint(1, 2147483647)})
                elif role == 4:
                    sql = "DELETE FROM Status WHERE ID_VK = " + str(id) + ";"
                    cursorR.execute(sql)
                    sql = "UPDATE Users SET Groupp = 7777, role = 4 WHERE ID_VK = " + str(id) + ";"
                    cursor.execute(sql)
                    conn.commit()

                    vk.method("messages.send", {"peer_id": id, "message": "Теперь я знаю о тебе достаточно). \n Используй кнопки клавиатуры.", "keyboard" : keyboards.getMainKeyboard(role = 4),
                                    "random_id": random.randint(1, 2147483647)})
                
                return False
            elif StatusR(id) == 1:
                if body.lower() == "справка":
                    vk.method("messages.send", {"peer_id": id, "message": "По вопросам сотрудничества писать на почту mr.woodysimpson@gmail.com \n Чтобы позвать администратора, введите Позвать.", "keyboard" : keyboards.keyboardRef1, 
                                    "random_id": random.randint(1, 2147483647)})
                    sql = "UPDATE Status SET Status = 15 WHERE ID_VK = " + str(id) + ";"
                    cursorR.execute(sql)
                    conn.commit()
                    return False
                elif body.lower() == "продолжить регистрацию":
                    vk.method("messages.send", {"peer_id": id, "message": "Введите свое имя в чат", "keyboard" : "", 
                                    "random_id": random.randint(1, 2147483647)})
                    sql = "UPDATE Status SET Status = 1 WHERE ID_VK = " + str(id) + ";"
                    cursorR.execute(sql)
                    conn.commit()
                    return False
                today = datetime.date.today()
                sql = "UPDATE users SET name = '" + str(body) + "' WHERE ID_VK = " + str(id)
                cursor.execute(sql)
                
                sql = "UPDATE Status SET Status = 2 WHERE ID_VK = " + str(id) + ";"
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send", {"peer_id": id, "message": "Очень приятно, " + str(body) + "",
                                    "random_id": random.randint(1, 2147483647)})
                vk.method("messages.send", {"peer_id": id, "message": "Расписание какой группы мне тебе показывать?\n Отправь сообщение с номером твоей группы.", "keyboard" : keyboards.get_undo,
                                    "random_id": random.randint(1, 2147483647)})
                return False
            elif StatusR(id) == 2:
                try:
                    if showGroupId(body):
                        if int(body) > 1100 and int(body)<10000:
                            sql = "UPDATE Users SET Groupp= " + str(showGroupId(body)) + " ,groupReal = " + str(body)+ " WHERE ID_VK = " + str(id) + ";"
                            cursor.execute(sql)
                            connection.commit()
                            sql = "DELETE FROM Status WHERE ID_VK = " + str(id) + ";"
                            cursorR.execute(sql)
                            conn.commit()
                            UserParams.update(int(MessageSettings.id))
                            vk.method("messages.send", {"peer_id": id, "message": "Твоя группа: " + body + "\n Теперь мне все понятно и ты можешь пользоваться ботом :)\n Настоятельно рекомендую подписаться на оффициальную группу @botraspisanie. Здесь ты сможешь получить много полезной информации.",
                                                        "keyboard" : keyboards.keyboardInfo,
                                                "random_id": random.randint(1, 2147483647)})
                            vk.method("messages.send",
                                      {"peer_id": id,
                                       "message": "Меню"
                                          , "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                       "random_id": random.randint(1, 2147483647)})
                            # vk.method("messages.send",
                            #         {"peer_id": id, "sticker_id" : 6880 , "random_id": random.randint(1, 2147483647)})
                            # vk.method("messages.send", {"peer_id": id, "random_id": random.randint(1, 2147483647), "attachment": "poll-182372147_348171795"})

                        elif int(body) > 10000:
                            sql = "UPDATE Users SET Groupp= " + str(showGroupId(body)) + " ,groupReal = " + str(body) + ", affiliate = true, role = 6 WHERE ID_VK = " + str(id) + ";"
                            cursor.execute(sql)
                            connection.commit()
                            sql = "DELETE FROM Status WHERE ID_VK = " + str(id) + ";"
                            cursorR.execute(sql)
                            conn.commit()
                            UserParams.update(int(MessageSettings.id))
                            vk.method("messages.send",
                                {"peer_id": id, "message": "Ваше расписание может отсутствовать на сайте КАИ, однако вы"
                                       " можете добавить его самостоятельно, если его все же нет. Следуйте инструкциям!" # \n &#128295; Раздел в стадии тестирования :) &#128295;"
                                                           ,"keyboard" : keyboards.getMainKeyboard(6), "random_id": random.randint(1, 2147483647)})
                            vk.method("messages.send",
                                      {"peer_id": id, "message": "Инструкция приложена к этому сообщению.",
                                       "keyboard": keyboards.help_starosta_affiliate,
                                       "random_id": random.randint(1, 2147483647)})
                        else:
                            vk.method("messages.send",
                                      {"peer_id": id, "message": "Я не могу обработать такой номер группы. ",
                                       "keyboard" : keyboards.get_undo,
                                        "random_id": random.randint(1, 2147483647)})
                        return False
                    else:
                        vk.method("messages.send", {"peer_id": id, "message": "Что-что, а это точно не номер группы. Повтори ввод.",  "keyboard" : keyboards.get_undo,
                                                    "random_id": random.randint(1, 2147483647)})
                                
                except Exception as E:
                    print('Ошибка:\n', traceback.format_exc())
                    return False
            elif StatusR(id) == 4:
                try:

                    body = body.lower()
                    response = requests.post( BASE_URL_STAFF, data = "prepodLogin=" + str(body), headers = {'Content-Type': "application/x-www-form-urlencoded"}, params = {"p_p_id":"pubLecturerSchedule_WAR_publicLecturerSchedule10","p_p_lifecycle":"2","p_p_resource_id":"schedule"} )
                    # print(response.json())
                    if not len(response.json()):
                        vk.method("messages.send", {"peer_id": id, "message": "Расписание для вас отсутствует на сайте. Повторите ввод."
                                                                              " Возможно логин введен неверно.", "keyboard" : keyboards.get_undo,
                                                    "random_id": random.randint(1, 2147483647)})
                    else:
                        sql = "UPDATE Status SET status = 5 WHERE id_vk = " + str(id) + ";"
                        cursorR.execute(sql)
                        conn.commit()
                        sql = "UPDATE users SET login = '{}' WHERE id_vk ={}".format(body.lower(), id)
                        cursor.execute(sql)
                        connection.commit()
                        vk.method("messages.send", {"peer_id": id, "message": "Введите свою фамилию.", "keyboard": keyboards.get_undo,
                                                    "random_id": random.randint(1, 2147483647)})
                        return
                except Exception as E:
                    vk.method("messages.send",
                              {"peer_id": id, "message": "Расписание для вас отсутствует на сайте. Повторите ввод.",
                               "keyboard": keyboards.get_undo,
                               "random_id": random.randint(1, 2147483647)})
                    # print('Ошибка:\n', traceback.format_exc())
                    return False
            elif StatusR(id) == 5:
                try:

                    body = body.lower()
                    # print(str(body))
                    response = requests.post(
                        'https://kai.ru/for-staff/raspisanie?p_p_id=pubLecturerSchedule_WAR_publicLecturerSchedule10&p_p'
                        '_lifecycle=2&p_p_state=normal&p_p_mode=view&p_p_resource_id=getLecturersURL&p_p_cacheability='
                        'cacheLevelPage&p_p_col_id=column-1&p_p_col_count=1&query='+body)


                    if not len(response.json()):
                        vk.method("messages.send",
                                  {"peer_id": id, "message": "Расписание для вас отсутствует на сайте. Повторите ввод фамилии.",
                                   "keyboard": keyboards.get_undo,
                                   "random_id": random.randint(1, 2147483647)})
                    else:

                        sql = "SELECT login FROM users WHERE id_vk={}".format(id)
                        cursor.execute(sql)
                        login = cursor.fetchone()[0]
                        # print(login)
                        name = ""
                        for row in response.json():
                            # print(row,row["id"])
                            if row["id"].rstrip().lower() == login.rstrip().lower():
                                name = row["lecturer"]
                                sql = "DELETE FROM Status WHERE ID_VK = " + str(id) + ";"
                                cursorR.execute(sql)
                                conn.commit()
                                sql = "UPDATE users SET name='{}', role = 2 WHERE ID_VK = {}".format(name, id)
                                cursor.execute(sql)

                                connection.commit()
                                vk.method("messages.send", {"peer_id": id, "message": "Регистрация успешно завершена.",
                                                            "keyboard": keyboards.getMainKeyboard(2),
                                                            "random_id": random.randint(1, 2147483647)})
                        if not name:
                            vk.method("messages.send",
                                      {"peer_id": id,
                                       "message": "Совпадения не найдены. Повторите ввод. \nВведите логин без лишних символов.",
                                       "keyboard": keyboards.get_undo,
                                       "random_id": random.randint(1, 2147483647)})
                            sql = "UPDATE Status SET Status = 4 WHERE ID_VK = " + str(id) + ";"
                            cursorR.execute(sql)
                            conn.commit()
                            return

                except Exception as E:
                    print('Ошибка:\n', traceback.format_exc())
                    return False





            elif StatusR(id) == 15:
                try:
                    if body.lower() == "позвать":
                        vk.method("messages.send", {"peer_id": 159773942, "message": "Пользователь @id"+str(id) + " просит помощи" , "keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
                        vk.method("messages.send", {"peer_id": id, "message": "Сообщение администратору отправлено", "keyboard": keyboards.keyboardRef1, "random_id": random.randint(1, 2147483647)})
                        return False
                    elif body.lower() == "справка":
                        vk.method("messages.send", {"peer_id": id, "message": "По вопросам сотрудничества писать на почту mr.woodysimpson@gmail.com \n Чтобы позвать администратора, введите Позвать.", "keyboard" : keyboards.keyboardRef, 
                                        "random_id": random.randint(1, 2147483647)})
                        sql = "UPDATE Status SET Status = 15 WHERE ID_VK = " + str(id) + ";"
                        cursor.execute(sql)
                        conn.commit()
                        return False
                    elif body.lower() == "продолжить регистрацию":
                        vk.method("messages.send", {"peer_id": id, "message": "Введите свое имя в чат", "keyboard" : keyboards.get_undo, 
                                        "random_id": random.randint(1, 2147483647)})
                        sql = "UPDATE Status SET Status = 1 WHERE ID_VK = " + str(id) + ";"
                        cursorR.execute(sql)
                        conn.commit()
                        return False
                            
                                
                    return False
                except Exception as E:
                    vk.method("messages.send", {"peer_id": id, "message": "Error status 15 in reg", 
                                        "random_id": random.randint(1, 2147483647)})
                return False
    except SystemExit:
        sys.exit(1)
    except:  
        print('Ошибка:\n', traceback.format_exc())  



BASE_URL = 'https://kai.ru/raspisanie'
BASE_URL_STAFF = "https://kai.ru/for-staff/raspisanie"

def getGroupsResponse(groupNumber):
    try:
        cursor.execute("SELECT shedule FROM saved_timetable WHERE groupp = 1")
        result = cursor.fetchone()[0]
        result = json.loads(result)
        for elem in result:
            if int(elem["group"]) == int(groupNumber):

                return elem["id"]
        return False
    except:
        print('Ошибка GET GROUP RESPONSE:\n', traceback.format_exc())
        return False




def showGroupId(groupNumber):
    id = int(MessageSettings.id)
    try:
        response = requests.post( BASE_URL + "?p_p_id=pubStudentSchedule_WAR_publicStudentSchedule10&p_p_lifecycle=2&p_p_resource_id=getGroupsURL&query=" + groupNumber, headers = {'Content-Type': "application/x-www-form-urlencoded"}, params = {"p_p_id":"pubStudentSchedule_WAR_publicStudentSchedule10","p_p_lifecycle":"2","p_p_resource_id":"schedule"}, timeout = 4)
        if str(response.status_code) != '200':
            raise ConnectionError
            # vk.method("messages.send",
            #     {"peer_id": id, "message": "&#9888;Ошибка подключения к серверам.&#9888; \n Вероятно, на стороне kai.ru произошел сбой. Вам необходимо продолжить регистрацию как только сайт kai.ru станет доступным.", "random_id": random.randint(1, 2147483647)})
            # vk.method("messages.send",
            #         {"peer_id": id, "message": "test" , "sticker_id" : 18486 , "random_id": random.randint(1, 2147483647)})
            
            return False
        print(response.status_code, response.json(), response)
        response = response.json()[0]
        return response['id']
    except IndexError:
        vk.method("messages.send",
                {"peer_id": id, "message": "Такой группы нет.", "random_id": random.randint(1, 2147483647)})
        return False
    except (ConnectionError, TimeoutError, requests.exceptions.ReadTimeout):
        try:
            group = getGroupsResponse(groupNumber)
            if group:
                return group
            vk.method("messages.send",
                    {"peer_id": id, "message": "&#9888;Ошибка подключения к серверам.&#9888; \n Вероятно, на стороне kai.ru произошел сбой. Вам необходимо продолжить регистрацию (ввод номера группы) как только сайт kai.ru станет доступным.", "random_id": random.randint(1, 2147483647)})
            vk.method("messages.send",
                    {"peer_id": id, "message": "test" , "sticker_id" : 18486 , "random_id": random.randint(1, 2147483647)})
            return False
        except:
            print('Ошибка:\n', traceback.format_exc())
        return False
    except:
        print('Ошибка:\n', traceback.format_exc())
        return False






def InBaseR(id): ### Проверка на зарегестрированность и наличие в базе Status (RAM)
    sql = "SELECT Status FROM Status WHERE ID_VK=" + str(id) +";"
    cursorR.execute(sql)
    res=cursorR.fetchall()
    if len(res)==0:
        return True
    else:
        return False

def InBase(id): ### Проверка на зарегестрированность и наличие в базе Users
    try:
        #global allCommands, statUser
        #if MessageSettings.statUser.count(id)>1:
        #    MessageSettings.statUser.remove(id)
        #elif MessageSettings.statUser.count(id)==1:
        #    pass;
        #else:
        #    MessageSettings.statUser.append(id)
        #if statUser.count(id) == 0:
        #    statUser.append(id)
        #    MessageSettings.statUser = len(statUser)
        #    print(MessageSettings.statUser)
            
        #allCommands += 1
        #MessageSettings.allCommands = allCommands
        sql = "SELECT Groupp, login, name FROM Users WHERE ID_VK=" + str(id) + ";"
        cursor.execute(sql)
        res=cursor.fetchone()
        if res == None:
            return False
        group = res[0]

        login = ""
        try:
            login = res[1]
        except Exception as E:
            print('Ошибка:\n', traceback.format_exc())
        if login and res[2].lstrip().rstrip():
            return True

        if len(str(group)) == 0:
            return False
        elif int(group) == 0 and int(id)<2000000000:
            return False
        else:
            return True
    except TypeError:
        if int(id)>2000000000:
            return True
        return False
    except Exception as E:
        print('Ошибка:\n', traceback.format_exc())

        vk.method("messages.send", {"peer_id": id, "message": "Что-то пошло не так.", 
                                    "random_id": random.randint(1, 2147483647)})
        vk.method("messages.send",
                        {"peer_id": id, "message": "test" , "sticker_id" : 6890 , "random_id": random.randint(1, 2147483647)})
        vk.method("messages.send", {"peer_id": id, "message": "Перезагружаюсь...", 
                                    "random_id": random.randint(1, 2147483647)})
        print("GLOBAL ERROR - RESTART ")
        sys.exit(1)



def StatusR(id): ### Текущий статус в таблице Status (RAM)
    sql = "SELECT Status FROM Status WHERE ID_VK=" + str(id)
    cursorR.execute(sql)
    res=cursorR.fetchone()
    if res == None:
        return
    if int(res[0]) > 0:
        return int(res[0])
    else:
        return
    

@jit(nopython=True, fastmath=True, cache=True)
def damerau_levenshtein_distance(s1, s2):
    d = {}
    lenstr1 = len(s1)
    lenstr2 = len(s2)
    for i in range(-1, lenstr1 + 1):
        d[(i, -1)] = i + 1
    for j in range(-1, lenstr2 + 1):
        d[(-1, j)] = j + 1
    for i in range(lenstr1):
        for j in range(lenstr2):
            if s1[i] == s2[j]:
                cost = 0
            else:
                cost = 1
            d[(i, j)] = min(
                d[(i - 1, j)] + 1,  # deletion
                d[(i, j - 1)] + 1,  # insertion
                d[(i - 1, j - 1)] + cost,  # substitution
            )
            if i and j and s1[i] == s2[j - 1] and s1[i - 1] == s2[j]:
                d[(i, j)] = min(d[(i, j)], d[i - 2, j - 2] + cost)  # transposition
    return d[lenstr1 - 1, lenstr2 - 1]



def CheckStatus():
    body = MessageSettings.getText()
    id = MessageSettings.getId()
    button = MessageSettings.button
    try:
        today = datetime.date.today()
        body = MessageSettings.getText()
        status = UserParams.Status
        id = MessageSettings.getId()




        if body.lower() == "выход" or body.lower() == "назад" or body.lower() == "выйти":
            try:
                vk.method("messages.send",
                                {"peer_id": id, "message": "Главное меню", "keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
                cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
                cursorR.execute("DELETE FROM NoteR WHERE ID_VK="+str(id))
                cursorR.execute("DELETE FROM Task WHERE ID_VK="+str(id))
                cursor.execute("DELETE FROM storage WHERE ischecked = -1 and id_vk = " + str(id))
                cursorR.execute("DELETE FROM storage WHERE id="+str(id))
                conn.commit()
                connection.commit()

            except Exception as E:

                print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 46: # ADMIN DISTRIBUTION

            sql = "SELECT id_vk FROM users WHERE ID_VK < 2000000000 AND role = 1".format()
            cursor.execute(sql)
            result_users = cursor.fetchall()
            #pprint(result_users)

            code = """var message = '{}';
            var attachment = '{}';
            """.format(MessageSettings.getText(), MessageSettings.GetAttachments())
            # code += """var keyboard = "{'one_time': false, 'buttons': [[{'action': {'type': text', 'payload": '{'button': 'tomorrow'}', 'label': 'На завтра'}, 'color': 'primary'}, {'actio': {'type': 'text', 'payload': '{'button': 'exams'}, 'label': 'Экзамены'}, 'color': 'positive'}], [{"action": {"type": "text", "payload": "{\"button\": \"today\"}", "label": "На сегодня"}, "color": "primary"}, {"action": {"type": "text", "payload": "{\"button\": \"after\"}", "label": "На послезавтра"}, "color": "primary"}, {"action": {"type": "text", "payload": "{\"button\": \"all\"}", "label": "Полностью"}, "color": "primary"}], [{"action": {"type": "text", "payload": "{\"button\": \"chetnost\"}", "label": "Четность недели"}, "color": "default"}, {"action": {"type": "text", "payload": "{\"button\": \"task menu\"}", "label": "Задания и объявления"}, "color": "primary"}], [{"action": {"type": "text", "payload": "{\"button\": \"commands\"}", "label": "Команды"}, "color": "default"}, {"action": {"type": "text", "payload": "{\"button\": \"prepod\"}", "label": "Преподы"}, "color": "default"}], [{"action": {"type": "text", "payload": "\"{'button': 'feedback'}\"", "label": "Обратная связь"}, "color": "primary"}, {"action": {"type": "text", "payload": "{\"button\": \"profile\"}", "label": "Профиль"}, "color": "positive"}]]};"""
            current_list_users = []
            for each in result_users:
                current_list_users.append(each[0])
                if len(current_list_users) == 100:
                    users_string_join = ','.join(str(x) for x in current_list_users)
                    #print(users_string_join)
                    current_list_users = []
                    code += "API.messages.send({{ 'user_ids' : '{}', 'message' : message, 'attachment': attachment,'random_id' : {} }}); \n".format(users_string_join, random.randint(1, 2147483647) )
                    
            code += "return 100;"
            #print(code)
            cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
            conn.commit()
            print(vk.method("execute", {"code": code}))

            # message = "📩 Сообщение от старосты:\n" + MessageSettings.getText()
            # print(','.join(str(x[0]) for x in result_users))
            # vk.method("messages.send", {"user_ids": ','.join(str(x[0]) for x in result_users), "message": message,"attachment": MessageSettings.GetAttachments() ,"random_id": random.randint(1, 2147483647)})

            return "ok"
        elif status == 47:
            
            sql = "SELECT id_vk FROM users WHERE groupp = {} AND ID_VK < 2000000000 LIMIT 100".format(UserParams.groupId)
            cursor.execute(sql)
            result_users = cursor.fetchall()
            print(result_users)
            message = "📩 Сообщение от старосты:\n" + MessageSettings.getText()
            print(','.join(str(x[0]) for x in result_users))
            vk.method("messages.send", {"user_ids": ','.join(str(x[0]) for x in result_users), "message": message,"attachment": MessageSettings.GetAttachments() ,"random_id": random.randint(1, 2147483647)})
            cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
            conn.commit()
            return "ok"
        elif status == 48:
            try:
                domain  = body.partition("vk.com/")
                print(domain)
                if domain[1] == "vk.com/":
                    domain_id = domain[2]
                elif not domain[1] and not domain[2] and domain[0]:
                    domain_id = domain[0]
                else:
                    domain_id = False
                    vk.method("messages.send", {"peer_id": id, "message": "Некорректно. Повтори ввод",
                                                        "random_id": random.randint(1, 2147483647)})
                resp = vk.method("users.get", {"user_ids": str(domain_id)})
                print(resp)
                id_student = 0
                try:
                    id_student = resp[0]["id"]
                except KeyError:
                    vk.method("messages.send", {"peer_id": id, "message": "Ошибка. Такого человека нет.",
                                                        "random_id": random.randint(1, 2147483647)})
                sql = "SELECT * FROM users WHERE id_vk = {}".format(id_student)
                cursor.execute(sql)
                res = cursor.fetchone()
                print(res)
                if res:
                    student_groupId = int(res[2])

                else:
                    vk.method("messages.send", {"peer_id": id, "message": "Ошибка. Пользователь не зарегистрирован.",
                                                        "random_id": random.randint(1, 2147483647)})
                    return "ok"
                if UserParams.adminLevel >= 2:
                    if UserParams.groupId != student_groupId:
                        vk.method("messages.send", {"peer_id": id, "message": "Ошибка. Пользователь не из вашей группы",
                                                        "random_id": random.randint(1, 2147483647)})
                        return "ok"
                    sql = "UPDATE users SET groupreal = 0, groupp = 9999 WHERE ID_VK = {}".format( id_student )
                    cursor.execute(sql)
                    connection.commit()
                    vk.method("messages.send", {"peer_id": id, "message": "@id{} (Пользователь) был кикнут из вашей группы.".format(id_student),
                                    "random_id": random.randint(1, 2147483647)})
                    vk.method("messages.send", {"peer_id": id_student, "message": "Вы были кикнуты старостой из группы. Ваши настройки группы сброшены и это значит, что пока вы не установите свою группу в профиле, расписание будет недоступно","keyboard": keyboards.getMainKeyboard(1),
                                    "random_id": random.randint(1, 2147483647)})


                cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
                conn.commit()
                return "ok"
            except:
                vk.method("messages.send", {"peer_id": id, "message": "Некорректно. Повтори ввод",
                                            "random_id": random.randint(1, 2147483647)})

        elif status == 49:
            
            domain  = body.partition("vk.com/")
            print(domain)
            if domain[1] == "vk.com/":
                domain_id = domain[2]
            elif not domain[1] and not domain[2] and domain[0]:
                domain_id = domain[0]
            else:
                domain_id = False
                vk.method("messages.send", {"peer_id": id, "message": "Некорректно. Повтори ввод",
                                                    "random_id": random.randint(1, 2147483647)})
            resp = vk.method("users.get", {"user_ids": str(domain_id)})
            print(resp)
            id_student = 0
            try:
                id_student = resp[0]["id"]
            except KeyError:
                vk.method("messages.send", {"peer_id": id, "message": "Ошибка. Такого человека нет.",
                                                    "random_id": random.randint(1, 2147483647)})
            sql = "SELECT * FROM users WHERE id_vk = {}".format(id_student)
            cursor.execute(sql)
            res = cursor.fetchone()
            print(res)
            if res:
                student_groupId = int(res[2])
                student_warn_count = int(res[9])
            else:
                vk.method("messages.send", {"peer_id": id, "message": "Ошибка. Пользователь не зарегистрирован.",
                                                    "random_id": random.randint(1, 2147483647)})
                return "ok"
            if UserParams.adminLevel >= 2:
                if UserParams.groupId != student_groupId:
                    vk.method("messages.send", {"peer_id": id, "message": "Ошибка. Пользователь не из вашей группы",
                                                    "random_id": random.randint(1, 2147483647)})
                    return "ok"
                if student_warn_count >= 2:
                    sql = "UPDATE users SET warn = {}, expiration = '{}', role = 5 WHERE ID_VK = {}".format(student_warn_count + 1, datetime.date(today.year, today.month, today.day) + datetime.timedelta(days = 61), id_student )
                    cursor.execute(sql)
                    connection.commit()
                    vk.method("messages.send", {"peer_id": id, "message": "@id{} (Пользователь) был заблокирован на 2 месяца".format(id_student),
                                    "random_id": random.randint(1, 2147483647)})
                    vk.method("messages.send", {"peer_id": id_student, "message": "Вы были заблокированы на 2 месяца за нарушение правил.","keyboard": keyboards.warnList,
                                    "random_id": random.randint(1, 2147483647)})
                else:
                    sql = "UPDATE users SET warn = {}, expiration = '{}' WHERE ID_VK = {}".format(student_warn_count + 1, datetime.date(today.year, today.month, today.day) + datetime.timedelta(days = 61), id_student )
                    cursor.execute(sql)
                    connection.commit()
                    vk.method("messages.send", {"peer_id": id, "message": "@id{} (Пользователь) получил предупреждение".format(id_student),
                                    "random_id": random.randint(1, 2147483647)})
                    vk.method("messages.send", {"peer_id": id_student, "message": "Вам выдано предупреждение за нарушение правил.", "keyboard": keyboards.warnList,
                                    "random_id": random.randint(1, 2147483647)})     
            
            cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
            conn.commit()
            return "ok"

        elif status == 50:
            date = str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5))
            try:
                try:
                    #print(body[:2])
                    
                    if ((int)(body[:2]) and (int)(body[3:]) and body[2] == "." and (int)(body[:2])<32 and (int)(body[3:])<13):
                        date = str(datetime.datetime.now().year) + "-" + body[3:] + "-" + body[:2]
                               
                    else:
                        pass
                        #vk.method("messages.send", {"peer_id": id, "message": "Формат неверный!", 
                        #                "random_id": random.randint(1, 2147483647)})
                        return "ok"
                        #print(111)
                            
                except Exception as E:
                    return "ok"

                       
            
            except Exception as E:
                vk.method("messages.send", {"peer_id": id, "message": "Формат некорректный. Верный формат - 'дд.мм' ", 
                                    "random_id": random.randint(1, 2147483647)})
            finally:
                
                if date == str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5)) and body != "Через неделю" and body != "Через 2 недели":
                    vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                    "random_id": random.randint(1, 2147483647)})
                    return "ok"
                        
                if body == "Через неделю":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=7))
                elif body == "Через 2 недели":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=14))
                        
                else:
                    try:
                        #print(datetime.date(int(date[0:4]), int(date[5:7]), int(date[8:])) + datetime.timedelta(days=30))
                        if datetime.date(int(date[0:4]), int(date[5:7]), int(date[8:])) > datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=30):
                            vk.method("messages.send", {"peer_id": id, "message": "Запланированная дата неверная. \n Разрешено добавлять задания только в течение следующего месяца", "keyboard": keyboards.keyboardAddTasks2,
                                            "random_id": random.randint(1, 2147483647)})
                            return "ok"
                        elif date == str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5)):
                            vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                                        "random_id": random.randint(1, 2147483647)})
                            return "ok"
                    except Exception as E:
                        print('Ошибка:\n', traceback.format_exc())



                sql = "INSERT INTO Task VALUES(" + str(id) + ", '" + date + "')"
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send", {"peer_id": id, "message": "Введите задание и к этому же сообщению прикрепите медиавложение (фото/видео/аудио/документ)", "keyboard": keyboards.keyboardAddTasks2,
                                                "random_id": random.randint(1, 2147483647)})
                sql = "UPDATE Status SET Status = 51 WHERE ID_VK = " + str(id) 
                cursorR.execute(sql)
                conn.commit()
                return "ok"
        elif status == 51:
            id = MessageSettings.getId()
            level = UserParams.adminLevel
            sql="SELECT COUNT(*) FROM Task WHERE UserID = " + str(id)
            cursor.execute(sql)
            try:
                count = (int)(cursor.fetchone()[0])
            except Exception as E:
                count = 0
            #print(count)
            if (count > 10 and level < 2):
                vk.method("messages.send", {"peer_id": id, "message": "Превышено допустимое число активных заданий. Ваш лимит: 10" , "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                                "random_id": random.randint(1, 2147483647)})
                cursor.execute(sql)
                sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                cursorR.execute(sql)
                sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
                cursorR.execute(sql)
                connection.commit()
                conn.commit()


                return "ok"
            sql = "SELECT Datee FROM Task WHERE ID_VK = " +str(id)
            cursorR.execute(sql)
            date = cursorR.fetchone()
            date = str(date)[2:-3]
            #print("DATA--------------------- " + str(date))
            sql = "SELECT MAX(ID) FROM Task"
            cursor.execute(sql)
            res = cursor.fetchone()[0]
            res = res if res != None else 0
            # count = (int)(str(res)[1:-2]) + 1
            # count = (int)(str(cursor.fetchone())[1:-2]) + 1
            count = int(res)+1

            user_info = """{{"type" : "message","owner_id" : {},"peer_id": {},"conversation_message_id" : {}}}""".format(id, id, MessageSettings.messageId)



            # sql = "INSERT INTO Task VALUES (" + str(count) + ", " + str(UserParams.groupId) + ", " + str(id) + ", '" + str(date) + "', '" + str(MessageSettings.getText()) + "', '" + str(MessageSettings.GetAttachments()) + "', 0)"
            sql = "INSERT INTO Task VALUES ({count}, {group_id}, {id}, '{date}', '{text}', '{attachments}', 0, '{user_info}')".format(
                count = count, group_id = UserParams.groupId, id = id, date = date, text = MessageSettings.getText(), attachments = MessageSettings.GetAttachments(), user_info = user_info
            )
            cursor.execute(sql)
            #print(sql)
            sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            connection.commit()
            conn.commit()

            vk.method("messages.send", {"peer_id": id, "message": "Задание успешно добавлено на " + date, "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                                "random_id": random.randint(1, 2147483647)})
            return "ok"
        elif status == 52:
            date = str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5))
            try:
                try:
                    #print(body[:2])
                    print(body[:2], body[3])
                    if ((int)(body[:2]) and (int)(body[3:]) and body[2] == "." and (int)(body[:2])<32 and (int)(body[3:])<13):
                        date = str(datetime.datetime.now().year) + "-" + body[3:] + "-" + body[:2]
                               
                    else:
                        pass
                        #vk.method("messages.send", {"peer_id": id, "message": "Формат неверный!", 
                        #                "random_id": random.randint(1, 2147483647)})
                        return "ok"
                        #print(111)
                            
                except Exception as E:
                    return "ok"

                       
            
            except Exception as E:
                vk.method("messages.send", {"peer_id": id, "message": "Формат некорректный. Верный формат - 'дд.мм' ", 
                                    "random_id": random.randint(1, 2147483647)})
            finally:
                
                if date == str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5)) and body != "Через неделю" and body != "Через 2 недели":
                    vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                    "random_id": random.randint(1, 2147483647)})
                    return "ok"
                        
                if body == "Через неделю":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=7))
                elif body == "Через 2 недели":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=14))
                        
                else:
                    try:
                        if datetime.date(int(date[0:4]), int(date[5:7]), int(date[8:])) > datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=30):
                            vk.method("messages.send", {"peer_id": id, "message": "Запланированная дата неверная. \n Разрешено добавлять объявления только в течение следующего месяца", "keyboard": keyboards.keyboardAddTasks2,
                                            "random_id": random.randint(1, 2147483647)})
                            return "ok"
                        elif date == str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5)):
                            vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                                        "random_id": random.randint(1, 2147483647)})
                            return "ok"
                    except Exception as E:
                        pass



                sql = "INSERT INTO Task VALUES(" + str(id) + ", '" + date + "')"
                #print(sql)
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send", {"peer_id": id, "message": "Введите текст объявления.", "keyboard": keyboards.keyboardAddTasks2,
                                                "random_id": random.randint(1, 2147483647)})
                sql = "UPDATE Status SET Status = 53 WHERE ID_VK = " + str(id) 
                cursorR.execute(sql)
                conn.commit()
                return "ok"
        elif status == 53:
            id = MessageSettings.getId()
            level = UserParams.adminLevel
            #print("Stage 1: " + str(id) + " " + str(level))
            sql='SELECT COUNT(*) FROM "Adv" WHERE userid = ' + str(id)
            try:
                sql = 'SELECT COUNT(*) FROM "Adv" WHERE userid = ' + str(id)
                cursor.execute(sql)
                try:
                    count = (int)(cursor.fetchone()[0])
                except Exception as E:
                    count = 0
            
            
                if len(MessageSettings.getText()) > 250:
                    vk.method("messages.send", {"peer_id": id, "message": "Превышена максимальная длина объявления" , "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                                    "random_id": random.randint(1, 2147483647)})
                    return "ok"
                #print(count)

                if (count > 10 and level < 2):
                    vk.method("messages.send", {"peer_id": id, "message": "Превышено допустимое число активных заданий. Ваш лимит: 10" , "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                                    "random_id": random.randint(1, 2147483647)})
                    cursor.execute(sql)
                    sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                    cursorR.execute(sql)
                    sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
                    cursorR.execute(sql)
                    connection.commit()
                    conn.commit()
                    return "ok"
                sql = "SELECT Datee FROM Task WHERE ID_VK = " +str(id)
                cursorR.execute(sql)
                date = cursorR.fetchone()
                date = str(date)[2:-3]
                #print(date)
                #print("DATA--------------------- " + str(date))
                sql = 'SELECT MAX(id) FROM "Adv"'
                cursor.execute(sql)
                count = (int)(str(cursor.fetchone())[1:-2]) + 1
                sql = 'DELETE FROM "Adv" WHERE date = ' + "'" + str(date) + "' AND groupid = " + str(UserParams.groupId)
                sql = 'INSERT INTO "Adv" VALUES (' + str(count) + ", " + str(UserParams.groupId) + ", " + str(id) + ", '" + str(date) + "', '" + str(MessageSettings.getText()) + "')"
                cursor.execute(sql)
            except Exception as E:
                print('Ошибка:\n', traceback.format_exc())
            #print(sql)
            sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            connection.commit()
            conn.commit()

            vk.method("messages.send", {"peer_id": id, "message": "Объявление успешно добавлено на " + date, "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                                "random_id": random.randint(1, 2147483647)})
            return "ok"
        elif status == 55:
            id = MessageSettings.getId()
            body = MessageSettings.getText()
            if len(body) > 35:
                vk.method("messages.send", {"peer_id": id, "message": "Длина имени не должна превышать 35 символов. Повторите ввод, либо введите Выход.", "keyboard": keyboards.KeyboardProfile(),
                                                "random_id": random.randint(1, 2147483647)})
                return "ok"
            sql="UPDATE Users SET Name='" + str(body) + "' WHERE ID_VK="+str(id)
            cursor.execute(sql)
            cursor.execute('UPDATE users SET ischeked = 0 WHERE ID_VK = ' + str(id))
            connection.commit()
            UserParams.update(id)
            vk.method("messages.send", {"peer_id": id, "message": "Имя успешно изменено на: " + str(body), "keyboard": keyboards.KeyboardProfile(),
                                                "random_id": random.randint(1, 2147483647)})
            cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
            conn.commit()
            return "ok"
        elif status == 56:
            today = datetime.date.today()
            date = str(datetime.date(today.year, today.month, today.day))
            id = MessageSettings.getId()
            body = MessageSettings.getText()
            try:
                #print(body)
                realgroup = body
                group = showGroupId(realgroup)
                
                if (int)(body) > 1000 and (int)(body) < 100000 and group:
                    group = str(group)
                    admlevel = UserParams.adminLevel if UserParams.adminLevel != 2 else 1
                    print("Adm level", admlevel) 
                    sql = "UPDATE users SET groupp = {}, groupreal = {}, \"dateChange\" = '{}', admlevel = {} WHERE ID_VK = {}".format(group, str(realgroup), date, admlevel, id)
                    cursor.execute(sql)
                    cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
                    conn.commit()
                    connection.commit()

                    vk.method("messages.send",
                        {"peer_id": id, "message": "Изменено", "keyboard" : keyboards.getMainKeyboard(UserParams.role),
                            "random_id": random.randint(1, 2147483647)})
                elif (int)(body) > 10000 and False:
                    vk.method("messages.send",
                        {"peer_id": id, "message": "Ваше расписание не поддерживается ввиду его отсутствия на сайте КНИТУ-КАИ. Если вы уверены, что расписание существует на сайте, напишите об этом в Обсуждениях @botraspisanie", "keyboard": keyboards.keyboardAddTasks2, "random_id": random.randint(1, 2147483647)})
                else:
                    vk.method("messages.send",
                        {"peer_id": id, "message": "Повторите ввод.", "keyboard": keyboards.keyboardAddTasks2, "random_id": random.randint(1, 2147483647)})

            except Exception as E:
                #pass
                print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 57:
            id = MessageSettings.getId()
            date = str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5))
            try:
                try:
                    #print(body[:2])
                    print(body[:2], body[3])
                    if ((int)(body[:2]) and (int)(body[3:]) and body[2] == "." and (int)(body[:2])<32 and (int)(body[3:])<13):
                        date = str(datetime.datetime.now().year) + "-" + body[3:] + "-" + body[:2]
                               
                    else:
                        pass
                        #vk.method("messages.send", {"peer_id": id, "message": "Формат неверный!", 
                        #                "random_id": random.randint(1, 2147483647)})
                        return "ok"
                        #print(111)
                            
                except Exception as E:
                    return "ok"

                       
            
            except Exception as E:
                vk.method("messages.send", {"peer_id": id, "message": "Формат некорректный. Верный формат - 'дд.мм' ", 
                                    "random_id": random.randint(1, 2147483647)})
            finally:
                
                if date == str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5)) and body != "Через неделю" and body != "Через 2 недели":
                    vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                    "random_id": random.randint(1, 2147483647)})
                    return "ok"
                        
                if body == "Через неделю":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=7))
                elif body == "Через 2 недели":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=14))
                        
                else:
                    try:
                        if date == str(datetime.date(today.year, today.month, today.day) -  datetime.timedelta(days=5)):
                            vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                                        "random_id": random.randint(1, 2147483647)})
                            return "ok"
                    except Exception as E:
                        pass





                sql = "DELETE FROM \"Adv\" WHERE date = '{}' AND groupid = {}".format(date, UserParams.groupId)
                pprint(sql)
                cursor.execute(sql)
                connection.commit()
                vk.method("messages.send", {"peer_id": id, "message": "Объявления на указанную дату удалены.", "keyboard": keyboards.KeyboardProfile(),
                                                "random_id": random.randint(1, 2147483647)})
                cursorR.execute("DELETE FROM Status WHERE ID_VK=" + str(id))
                conn.commit()
                return "ok"
            
            return "ok"
        elif status == 58:
            id = MessageSettings.getId()
            body = MessageSettings.getText()
            try:
            
                vk.method("messages.send",
                    {"peer_id": 159773942, "message": "from @id" + str(id) + "\n" + body , "keyboard": keyboards.GetButtonAnswer(id), "attachment": MessageSettings.GetAttachments(), "random_id": random.randint(1, 2147483647)})
                vk.method("messages.send",
                    {"peer_id": id, "message": "Вопрос отправлен админу." , "keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})
                cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
                conn.commit()

            except Exception as E:
                pass
                #print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 59:
            id = MessageSettings.getId()
            body = MessageSettings.getText()
            button = MessageSettings.button
            try:
            
                sql = "SELECT userId FROM answers WHERE id = " + str(id)
                cursorR.execute(sql)
                idUser = cursorR.fetchone()[0]
                vk.method("messages.send",
                    {"peer_id": idUser, "message": "Ответ администратора:\n" + body, "attachment": MessageSettings.GetAttachments(), "random_id": random.randint(1, 2147483647)})
                vk.method("messages.send",
                    {"peer_id": id, "message": "Ответ отправлен","keyboard": keyboards.getMainKeyboard(UserParams.role), "random_id": random.randint(1, 2147483647)})

                cursorR.execute("DELETE FROM Status WHERE ID_VK="+str(id))
                cursorR.execute("DELETE FROM answers WHERE id = " + str(id))
                conn.commit()

            except Exception as E:
                #pass
                print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 60:
            id = MessageSettings.getId()
            body = MessageSettings.getText()

            try:


                cursorR.execute("INSERT INTO chatListen VALUES (" + str(id) + ", " + str(MessageSettings.getText()) + ")")
                conn.commit()
                cursorR.execute("DELETE FROM Status WHERE ID_VK=" + str(id))
                conn.commit()
            except Exception as E:
                cursorR.execute("DELETE FROM Status WHERE ID_VK=" + str(id))
                conn.commit()
                print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 180:
            id_vk = MessageSettings.getId()
            body = MessageSettings.getText().lower()
            att = MessageSettings.GetAttachments()
            sql = 'SELECT MAX(id) FROM storage'
            cursor.execute(sql)
            count = (int)(cursor.fetchone()[0]) + 1
            if not att:
                vk.method("messages.send",
                          {"peer_id": id, "message": "Отправьте сообщение с файлом.",
                            "random_id": random.randint(1, 2147483647)})
                return "ok"
            # (id, subject, media_vk, header, description, IsCheked, ChekedBy, id, dateChek)
            try:
                cursor.execute('INSERT INTO storage VALUES (' + str(id_vk) + ",' ','" + str(att) + "', ' ', ' ', -1, 0," + str(count) + ", '2020-01-01')")
                connection.commit()
                sql = "UPDATE Status SET Status = 181 WHERE ID_VK = " + str(id_vk) + ";"
                cursorR.execute(sql)
                conn.commit()
                cursorR.execute("INSERT INTO storage VALUES (" + str(id_vk) + ", " + str(count) + ")")
                conn.commit()
                vk.method("messages.send",
                          {"peer_id": id, "message": "Файл получен. \n Введите полное название предмета в именительном падеже. (прим. Дискретная математика)",
                            "random_id": random.randint(1, 2147483647)})
            except Exception as E:

                print('Ошибка:\n', traceback.format_exc())
            return "ok"

        elif status == 181:
            id = MessageSettings.getId()
            body = MessageSettings.getText().lower()
            att = MessageSettings.GetAttachments()
            cursorR.execute("SELECT * FROM storage WHERE id = " + str(id))
            media_id = cursorR.fetchone()[1]
            if len(body) > 200:
                vk.method("messages.send",
                          {"peer_id": id, "message": "Превышена максимально допустимая длина\nПовторите ввод",
                           "random_id": random.randint(1, 2147483647)})
                return "ok"
            # (id, subject, media_vk, header, description, IsCheked, ChekedBy, dateChek)
            try:
                cursor.execute("UPDATE storage SET subject = '" + body  + "' WHERE id = '" + str(media_id) + "'")
                connection.commit()
                sql = "UPDATE Status SET Status = 182 WHERE ID_VK = " + str(id) + ";"
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send",
                          {"peer_id": id, "message": "Введите заголовок.\n(методичка *название*, лаба №ч, практика, лекция *название*, вопросы к экзамену и тп.)",
                           "random_id": random.randint(1, 2147483647)})
            except Exception as E:

                print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 182:
            id = MessageSettings.getId()
            body = MessageSettings.getText().lower()
            att = MessageSettings.GetAttachments()
            cursorR.execute("SELECT * FROM storage WHERE id = " + str(id))
            media_id = cursorR.fetchone()[1]
            if len(body) > 200:
                vk.method("messages.send",
                          {"peer_id": id, "message": "Превышена максимально допустимая длина\nПовторите ввод",
                           "random_id": random.randint(1, 2147483647)})
                return "ok"
            # (id, subject, media_vk, header, description, IsCheked, ChekedBy, dateChek)
            try:
                cursor.execute("UPDATE storage SET header = '" + body  + "' WHERE id = '" + str(media_id) + "'")
                connection.commit()
                sql = "UPDATE Status SET Status = 183 WHERE ID_VK = " + str(id) + ";"
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send",
                          {"peer_id": id, "message": "Введите описание. В описании можно указать через какаую программу следует открывать файлы(-ы), какой преподаватель и тп.\n Пожалуйста, оставьте больше ключевых слов, по которым ваш файл можно будет найти, желательно в именительном падеже. За красиво и правильно оформленное описание возможна увеличенная награда.",
                           "random_id": random.randint(1, 2147483647)})
            except Exception as E:

                print('Ошибка:\n', traceback.format_exc())
            return "ok"
        elif status == 183:
            id = MessageSettings.getId()
            body = MessageSettings.getText().lower()
            att = MessageSettings.GetAttachments()
            cursorR.execute("SELECT * FROM storage WHERE id = " + str(id))
            media_id = cursorR.fetchone()[1]
            if len(body) > 500:
                vk.method("messages.send",
                          {"peer_id": id, "message": "Превышена максимально допустимая длина\nПовторите ввод",
                           "random_id": random.randint(1, 2147483647)})
                return "ok"
            # (id, subject, media_vk, header, description, IsCheked, ChekedBy, dateChek)
            try:
                cursor.execute("UPDATE storage SET description = '" + body + "' WHERE id = '" + str(media_id) + "'")
                connection.commit()
                cursor.execute("UPDATE storage SET ischecked = 0 WHERE id = '" + str(media_id) + "'")
                connection.commit()
                sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                cursorR.execute(sql)
                conn.commit()
                sql = "DELETE FROM storage WHERE id = " + str(id)
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send",
                          {"peer_id": id, "message": "Файл отправлен на проверку. Как только его проверит администратор, вы получите вознаграждение, а другие люди смогут увидеть ваш файл.","keyboard": keyboards.getMainKeyboard(UserParams.role),
                           "random_id": random.randint(1, 2147483647)})
            except Exception as E:

                print('Ошибка:\n', traceback.format_exc())
            return "ok"

        elif status == 190:
            id = MessageSettings.getId()
            try:
                body = MessageSettings.getText().replace(",", '')
                body = MessageSettings.getText().replace(".", '')
                body = MessageSettings.getText().replace("?", '')
                body = MessageSettings.getText().replace("!", '')
                body = body.lower()
                parts = body.split(' ')
                sql = "SELECT * FROM storage WHERE ischecked = 1 AND ("
                for elem in parts:
                    sql += "subject LIKE '%" + str(elem) + "%' or "
                    sql += "header LIKE '%" + str(elem) + "%' or "
                    sql += "description LIKE '%" + str(elem) + "%' or "
                    
                sql = sql[:-3] + ")"

                cursor.execute(sql)
                msg = "Результат поиска: \n"
                ans = ""
                for row in cursor:
                    ans += "id " + str(row[7]) + " ~ " + str(row[3]) + " == " + str(row[1]) + "\n"
                if ans == "":
                    ans = "❕ Результаты не найдены, попробуйте уточнить запрос."
                vk.method("messages.send",
                          {"peer_id": id, "message": msg + ans[:4095],"keyboard": keyboards.storageMain,
                           "random_id": random.randint(1, 2147483647)})
            except Exception as E:

                print('Ошибка:\n', traceback.format_exc())


            sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            conn.commit()
            return "ok"
        elif status == 191:
            id = MessageSettings.getId()
            try:
                media = MessageSettings.getText()
                try:
                    if int(media) > 0:
                        pass
                    else:
                        vk.method("messages.send", {"peer_id": id, "message": "Неверный формат, повторите ввод, либо введите 'Выход'" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                except Exception:
                    #print("исключение")
                    try:
                        if media[:len("Скачать файл №")] == "Скачать файл №":
                            media = media[:len("Скачать файл №")]
                        else:
                            vk.method("messages.send", {"peer_id": id, "message": "Неверный формат, повторите ввод, либо введите 'Выход'" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                            return "ok"
                    except Exception:
                        #print("исключение2")
                        vk.method("messages.send", {"peer_id": id, "message": "Неверный формат, повторите ввод, либо введите 'Выход'" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                

                sql = "SELECT * FROM storage WHERE ischecked = 1 and id = " + str(media)
                cursor.execute(sql)
                res = cursor.fetchone()
                if res:
                    msg = "📒 id " +  str(res[7]) + " \n📝Предмет: " +  str(res[1]) + "\n📄Заголовок: " + str(res[3]) + "\n📂Описание: " + str(res[4])
                    vk.method("messages.send", {"peer_id": id, "message": msg ,"keyboard": keyboards.getdownloadstorage(res[7]), "random_id": random.randint(1, 2147483647)})
                else:
                    msg = "Файл не найден."
                    vk.method("messages.send", {"peer_id": id, "message": msg ,"keyboard": keyboards.storageMain, "random_id": random.randint(1, 2147483647)})
                
            except Exception as E:
                print('Ошибка:\n', traceback.format_exc())


            sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            conn.commit()
            return "ok"

        elif status == 199:
            Gamehub(id, vk, button)
            return "ok"

        elif status == 205: # Мафия, создание комнаты, выбор числа игроков
            try:
                if int(body) >= 2 and int(body) <=20:
                    sql = 'SELECT MAX(id) FROM mafiaRooms'
                    cursorR.execute(sql)
                    idRoom = 0
                    try:
                        idRoom = cursor.fetchone()[0]
                    except Exception:
                        pass
                    sql = "INSERT INTO mafiaRooms VALUES (" + str(int(idRoom)+1) + "," + str(body) + ",0,0)"
                    cursorR.execute(sql)
                    sql = "INSERT INTO mafiaUsers VALUES (1," + str(int(idRoom)+1) + "," + str(id) + ",0,0)"
                    cursorR.execute(sql)
                    sql = "UPDATE Status SET Status = 199 WHERE ID_VK = " + str(id) + ";"
                    cursorR.execute(sql)
                    conn.commit()
                    vk.method("messages.send", {"peer_id": id, "message": "[МАФИЯ]: Ожидаем подключения остальных игроков. Вы уже можете начать общаться в чате - просто отправь любое текстовое сообщение", "random_id": random.randint(1, 2147483647)})
                else:
                    vk.method("messages.send", {"peer_id": id, "message": "Выберите число участников от 3 до 20, либо введите 'Выход'" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
            except Exception:
                print('Ошибка:\n', traceback.format_exc())
                vk.method("messages.send", {"peer_id": id, "message": "Выберите число участников от 3 до 20, либо введите 'Выход'" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
            return "ok"
        elif status == 206: # подключение к комнате, запуск игры
            try:
                if int(body) >= -1:
                    sql = "SELECT * FROM mafiaRooms WHERE id = " + str(body)
                    cursorR.execute(sql)
                    res = cursorR.fetchone()
                    #print("res in main - ", res)
                    maxPlayers = 0
                    statusGame = -1
                    room = -1
                    if res:
                        maxPlayers = res[1]
                        statusGame = res[2]
                        room = res[0]
                    else:
                        vk.method("messages.send", {"peer_id": id, "message": "ID не найден. Повторите ввод, либо введите Выход" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    if statusGame == 0:
                        #print(room)
                        sql = "SELECT COUNT(*) FROM mafiaUsers WHERE room = " + str(room)
                        #print(sql)
                        cursorR.execute(sql)
                        countUsers = cursorR.fetchone()[0]
                        if countUsers < maxPlayers:
                            sql = "SELECT MAX(id) FROM mafiaUsers WHERE room = " + str(room)
                            cursorR.execute(sql)
                            maxId = cursorR.fetchone()[0]
                            currentId = int(maxId) + 1
                            sql = "INSERT INTO mafiaUsers VALUES (" + str(currentId) + "," + str(room) + "," + str(id) + ",0,0)"
                            cursorR.execute(sql)

                            
                            sql = "SELECT ID_VK FROM mafiaUsers WHERE room = " + str(room)
                            cursorR.execute(sql)
                            res = cursorR.fetchall()
                            ids = ''
                            for row in res:
                                ids += str(row[0]) + ","
                            vk.method("messages.send", {"user_ids": ids[:-1], "message": "@id" + str(id) + " (Игрок) подсоединился к игре","keyboard": keyboards.keyboardNull, "random_id": random.randint(1, 2147483647)})

                            sql = "UPDATE Status SET Status = 199 WHERE ID_VK = " + str(id) + ";"
                            #print(sql)
                            cursorR.execute(sql)
                            conn.commit()
                            if countUsers + 1 >= maxPlayers:
                                sql = "SELECT ID_VK FROM mafiaUsers WHERE room = " + str(room)
                                cursorR.execute(sql)
                                res = cursorR.fetchall()
                                ids = ''
                                for row in res:
                                    ids += str(row[0]) + ","
                                #print(ids)
                                sql = "UPDATE mafiaRooms SET status = 1 WHERE id = " + str(room)
                                cursorR.execute(sql)
                                mesg = "\n Первый день в этом городке. Вам предстоит ВЫЖИТЬ в этой схватке закона и криминала."
                                mesg += "\n [Подсказка] Вы можете поздороваться с другими игроками, как только вы будете готовы начать ход - нажмите на кнопку Готов"
                                vk.method("messages.send", {"user_ids": ids[:-1], "message": "[МАФИЯ] Команда собрана! Игроков участвует: " + str(maxPlayers) + mesg,"keyboard": keyboards.mafia_acceptgame, "random_id": random.randint(1, 2147483647)})

                            else:
                                vk.method("messages.send", {"peer_id": id, "message": "[МАФИЯ]: Ожидаем подключения остальных игроков. Вы уже можете начать общаться в чате - просто отправь любое текстовое сообщение", "random_id": random.randint(1, 2147483647)})

                    else:
                         vk.method("messages.send", {"peer_id": id, "message": "Комната переполнена. Выберите ID другой комнаты :(" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                else:
                    vk.method("messages.send", {"peer_id": id, "message": "Выберите ID комнаты :(" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
            except Exception:
                print('Ошибка:\n', traceback.format_exc())
                vk.method("messages.send", {"peer_id": id, "message": "Выберите ID комнаты :)" ,"keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
            return "ok"

        elif status == 301:
            try:
                id = MessageSettings.getId()
                response = requests.post(BASE_URL_STAFF, data="prepodLogin=" + str(UserParams.login),
                                         headers={'Content-Type': "application/x-www-form-urlencoded"},
                                         params={"p_p_id": "pubLecturerSchedule_WAR_publicLecturerSchedule10",
                                                 "p_p_lifecycle": "2", "p_p_resource_id": "schedule"})
                if str(response.status_code) != '200':
                    vk.method("messages.send", {"peer_id": id, "message": "&#9888; Не удалось запросить список ваших групп. Возникла ошибка при подключении к серверам. \nКод ошибки: {0} &#9888;".format(
                        str(response.status_code)),
                                                "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                    cursorR.execute(sql)
                    conn.commit()
                    return "ok"
                response = response.json()
                if len(response) == 0:
                    vk.method("messages.send", {"peer_id": id,
                                                "message": "&#9888;Не удалось запросить список ваших групп. Расписание пустое.&#9888;",
                                                "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                    cursorR.execute(sql)
                    conn.commit()
                    return "ok"
                groups = set()

                for day in response.keys():
                    for item in response[day]:
                        groups.add(item["group"])
                try:
                    body = int(body)

                    # body = showGroupId(body)
                    # assert not body
                    if str(body) not in groups:
                        vk.method("messages.send", {"peer_id": id,
                                                    "message": "&#9888;Вы не преподаете у данной группы \n Введите другой номер группы или нажмите Выход.!&#9888;",
                                                    "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    else:
                        vk.method("messages.send", {"peer_id": id,
                                                    "message": "Введите сообщение студентам. К сообщению можно прикрепить медиавложения.",
                                                    "keyboard": keyboards.exit,
                                                    "random_id": random.randint(1, 2147483647)})
                        sql = "INSERT INTO prepod_users VALUES ({},{}, 0)".format(id, body)
                        cursorR.execute(sql)
                        sql = "UPDATE status SET status = 302 WHERE id_vk = {}".format(id)
                        cursorR.execute(sql)
                        conn.commit()
                        return "ok"
                except:
                    print('Ошибка:\n', traceback.format_exc())
                    vk.method("messages.send", {"peer_id": id,
                                                "message": "&#9888;Введите корректный номер группы!&#9888;",
                                                "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
            except Exception:
                print('Ошибка:\n', traceback.format_exc())

            return "ok"
        elif status == 302:
            sql = "SELECT * FROM prepod_users WHERE id = {}".format(id)
            cursorR.execute(sql)
            groupId = cursorR.fetchone()[1]
            sql = "SELECT id_vk FROM users WHERE groupreal = {} AND ID_VK < 2000000000 LIMIT 100".format(groupId)
            cursor.execute(sql)
            result_users = cursor.fetchall()

            message = """📩 Сообщение от вашего преподавателя\n{}:\n
            {}""".format(UserParams.name, body)
            try:
                vk.method("messages.send", {"user_ids": ','.join(str(x[0]) for x in result_users), "message": message,
                                            "attachment": MessageSettings.GetAttachments(),
                                            "random_id": random.randint(1, 2147483647)})
            except:
                vk.method("messages.send", {"peer_id": id,
                                            "message": "Пользователи данной группы не зарегистрированы или сообщение пустое!".format(len(result_users)),
                                            "keyboard": keyboards.getMainKeyboard(2),
                                            "random_id": random.randint(1, 2147483647)})
                cursorR.execute("DELETE FROM Status WHERE ID_VK=" + str(id))
                conn.commit()
                sql = "DELETE FROM prepod_users WHERE id_vk = {}".format(id)
                cursorR.execute(sql)
                conn.commit()
                return "ok"

            cursorR.execute("DELETE FROM Status WHERE ID_VK=" + str(id))
            conn.commit()
            vk.method("messages.send", {"peer_id": id,
                                        "message": "Сообщение отправлено {} пользователям группы {}".format(len(result_users),groupId),
                                        "keyboard": keyboards.getMainKeyboard(2), "random_id": random.randint(1, 2147483647)})
            sql = "DELETE FROM prepod_users WHERE id = {}".format(id)
            cursorR.execute(sql)
            conn.commit()

            return "ok"
        elif status == 304:
            try:
                id = MessageSettings.getId()
                response = requests.post(BASE_URL_STAFF, data="prepodLogin=" + str(UserParams.login),
                                         headers={'Content-Type': "application/x-www-form-urlencoded"},
                                         params={"p_p_id": "pubLecturerSchedule_WAR_publicLecturerSchedule10",
                                                 "p_p_lifecycle": "2", "p_p_resource_id": "schedule"})
                if str(response.status_code) != '200':
                    vk.method("messages.send", {"peer_id": id, "message": "&#9888; Не удалось запросить список ваших групп. Возникла ошибка при подключении к серверам. \nКод ошибки: {0} &#9888;".format(
                        str(response.status_code)),
                                                "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                    cursorR.execute(sql)
                    conn.commit()
                    return "ok"
                response = response.json()
                if len(response) == 0:
                    vk.method("messages.send", {"peer_id": id,
                                                "message": "&#9888;Не удалось запросить список ваших групп. Расписание пустое.&#9888;",
                                                "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                    cursorR.execute(sql)
                    conn.commit()
                    return "ok"
                groups = set()
                groupId = 0
                for day in response.keys():
                    for item in response[day]:
                        groups.add(item["group"])
                        if item["group"] == str(body):
                            groupId = item["group"]
                try:
                    body = int(body)

                    # body = showGroupId(body)
                    # assert not body
                    if str(body) not in groups:
                        vk.method("messages.send", {"peer_id": id,
                                                    "message": "&#9888;Вы не преподаете у данной группы \n Введите другой номер группы или нажмите Выход.!&#9888;",
                                                    "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
                    else:
                        vk.method("messages.send", {"peer_id": id,
                                                    "message": "Введите запланированную дату задания для студентов.",
                                                    "keyboard": keyboards.keyboardAddTasks,
                                                    "random_id": random.randint(1, 2147483647)})
                        sql = "INSERT INTO prepod_users VALUES ({},{},{})".format(id, body, groupId)
                        cursorR.execute(sql)
                        sql = "UPDATE status SET status = 305 WHERE id_vk = {}".format(id)
                        cursorR.execute(sql)
                        conn.commit()
                        return "ok"
                except:
                    print('Ошибка:\n', traceback.format_exc())
                    vk.method("messages.send", {"peer_id": id,
                                                "message": "&#9888;Введите корректный номер группы!&#9888;",
                                                "keyboard": keyboards.exit, "random_id": random.randint(1, 2147483647)})
            except Exception:
                print('Ошибка:\n', traceback.format_exc())

            return "ok"
        elif status == 305:
            date = str(datetime.date(today.year, today.month, today.day) - datetime.timedelta(days=5))
            try:
                try:
                    # print(body[:2])

                    if ((int)(body[:2]) and (int)(body[3:]) and body[2] == "." and (int)(body[:2]) < 32 and (int)(
                            body[3:]) < 13):
                        date = str(datetime.datetime.now().year) + "-" + body[3:] + "-" + body[:2]

                    else:
                        pass
                        # vk.method("messages.send", {"peer_id": id, "message": "Формат неверный!",
                        #                "random_id": random.randint(1, 2147483647)})
                        return "ok"
                        # print(111)

                except Exception as E:
                    return "ok"



            except Exception as E:
                vk.method("messages.send", {"peer_id": id, "message": "Формат некорректный. Верный формат - 'дд.мм' ",
                                            "random_id": random.randint(1, 2147483647)})
            finally:

                if date == str(datetime.date(today.year, today.month, today.day) - datetime.timedelta(
                        days=5)) and body != "Через неделю" and body != "Через 2 недели":
                    vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                                "random_id": random.randint(1, 2147483647)})
                    return "ok"

                if body == "Через неделю":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=7))
                elif body == "Через 2 недели":
                    date = str(datetime.date(today.year, today.month, today.day) + datetime.timedelta(days=14))

                else:
                    try:
                        # print(datetime.date(int(date[0:4]), int(date[5:7]), int(date[8:])) + datetime.timedelta(days=30))
                        if datetime.date(int(date[0:4]), int(date[5:7]), int(date[8:])) > datetime.date(today.year,
                                                                                                        today.month,
                                                                                                        today.day) + datetime.timedelta(
                                days=30):
                            vk.method("messages.send", {"peer_id": id,
                                                        "message": "Запланированная дата неверная. \n Разрешено добавлять задания только в течение следующего месяца",
                                                        "keyboard": keyboards.keyboardAddTasks2,
                                                        "random_id": random.randint(1, 2147483647)})
                            return "ok"
                        elif date == str(
                                datetime.date(today.year, today.month, today.day) - datetime.timedelta(days=5)):
                            vk.method("messages.send", {"peer_id": id, "message": "Формат неверный, повторите ввод",
                                                        "random_id": random.randint(1, 2147483647)})
                            return "ok"
                    except Exception as E:
                        print('Ошибка:\n', traceback.format_exc())

                sql = "INSERT INTO Task VALUES(" + str(id) + ", '" + date + "')"
                cursorR.execute(sql)
                conn.commit()
                vk.method("messages.send", {"peer_id": id,
                                            "message": "Введите задание и к этому же сообщению прикрепите медиавложение (фото/видео/аудио/документ)",
                                            "keyboard": keyboards.keyboardAddTasks2,
                                            "random_id": random.randint(1, 2147483647)})
                sql = "UPDATE Status SET Status = 306 WHERE ID_VK = " + str(id)
                cursorR.execute(sql)
                conn.commit()
                return "ok"
        elif status == 306:
            id = MessageSettings.getId()
            level = UserParams.adminLevel
            sql = "SELECT COUNT(*) FROM Task WHERE UserID = " + str(id)
            cursor.execute(sql)
            try:
                count = (int)(cursor.fetchone()[0])
            except Exception as E:
                count = 0
            # print(count)
            if (count > 10 and level < 2):
                vk.method("messages.send",
                          {"peer_id": id, "message": "Превышено допустимое число активных заданий. Ваш лимит: 10",
                           "keyboard": keyboards.getMainKeyboard(UserParams.role),
                           "random_id": random.randint(1, 2147483647)})
                cursor.execute(sql)
                sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
                cursorR.execute(sql)
                sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
                cursorR.execute(sql)
                connection.commit()
                conn.commit()

                return "ok"
            sql = "SELECT Datee FROM Task WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            date = cursorR.fetchone()
            date = str(date)[2:-3]
            # print("DATA--------------------- " + str(date))
            sql = "SELECT MAX(ID) FROM Task"
            cursor.execute(sql)
            res = cursor.fetchone()[0]
            res = res if res != None else 0
            # count = (int)(str(res)[1:-2]) + 1
            # count = (int)(str(cursor.fetchone())[1:-2]) + 1
            count = int(res) + 1

            sql = "SELECT * FROM prepod_users WHERE id = {}".format(id)
            cursorR.execute(sql)
            groupId = cursorR.fetchone()[2]
            user_info = """{{"type" : "message","owner_id" : {},"peer_id": {},"conversation_message_id" : {}}}""".format(
                id, id, MessageSettings.messageId)

            prefix = "\nОт преподавателя {}:\n".format(UserParams.name)
            sql = "INSERT INTO Task VALUES ({count}, {group_id}, {id}, '{date}', '{text}', '{attachments}', 0, '{user_info}')".format(
                count=count, group_id=groupId, id=id, date=date, text=prefix + MessageSettings.getText(),
                attachments=MessageSettings.GetAttachments(), user_info=user_info
            )
            cursor.execute(sql)
            # print(sql)
            sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            sql = "DELETE FROM prepod_users WHERE id = " + str(id)
            cursorR.execute(sql)
            connection.commit()
            conn.commit()

            vk.method("messages.send", {"peer_id": id, "message": "Задание успешно добавлено на " + date,
                                        "keyboard": keyboards.getMainKeyboard(UserParams.role),
                                        "random_id": random.randint(1, 2147483647)})
            return "ok"

        elif status == 307:

            # sql = "DELETE FROM Task WHERE ID_VK = " + str(id)
            # cursorR.execute(sql)
            # sql = "DELETE FROM prepod_users WHERE id = " + str(id)
            # cursorR.execute(sql)
            groupId = showGroupId(body)
            if not groupId:
                return "ok"
            sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
            cursorR.execute(sql)
            conn.commit()
            msg = "Запрос отправлен на обработку"
            msg_id = vk.method("messages.send",
                               {"peer_id": MessageSettings.id, "message": msg, "keyboard": keyboards.exit,
                                "random_id": random.randint(1, 2147483647)})
            i = 1
            response = requests.post(("https://kai.ru/infoClick/-/info/group?id={id}").format(id=groupId))
            soup = BeautifulSoup(response.text, 'lxml')

            # print(soup.find("ul", attrs={ "id" : "mylist"}))
            list_students = soup.find(id="p_p_id_infoClick_WAR_infoClick10_")
            students = []
            result = ""
            for tag in list_students.find_all("td"):
                if len(tag.text) > 6:
                    name = (tag.text.strip().replace("\n", "").replace(
                        "                                                                Староста",
                        " (🙋 Староста)")).split(" ")
                    try:
                        name = name[0] + " " + name[1][0].capitalize() + "." + name[2][0].capitalize() + "."
                    except:
                        name = name[0][:20] + ". "
                    students.append(name)
            att = ""
            try:
                att = GetDocShedule(UserParams.groupId, MessageSettings.getPeer_id(), int(body), students)
            except:
                pass
            vk.method("messages.edit",
                      {"peer_id": MessageSettings.id, "message_id": msg_id, "message": "Бланк посещения","keyboard": keyboards.getMainKeyboard(2),
                       'attachment': att})

            return "ok"
        connection.commit()
        conn.commit()



        # CHECKING ACTUALITY OF CURRENT GROUP ID BY YEAR
        # if (UserParams.dateChange.year < today.year and today.month > 8) or (UserParams.dateChange.year == UserParams.dateChange.year and UserParams.dateChange.month < 9 and today.month > 8)
        body_excepts = ['Профиль', 'Изменить', 'Назад', '!группа', 'группа', 'Обратная связь']
        if False and UserParams.dateChange < datetime.date(2020, 9, 1) and UserParams.role in [1,3] and body not in body_excepts or (body.isdigit() and UserParams.dateChange < datetime.date(2020, 9, 1) ):
            vk.method("messages.send",
                        {"peer_id": id, "message": "&#9888; Кажется, номер вашей группы установлен в прошлом учебном году. Это означает, что вы не сможете получить актуальное расписание вашей группы. Обновить номер группы можно в профиле, нажав на номер группы или с помощью команды !группа", 
                        "random_id": random.randint(1, 2147483647)})
        # END CHECK

        return "no"
    except Exception as E:
        sql = "DELETE FROM Status WHERE ID_VK = " + str(id)
        cursorR.execute(sql)

        print('Ошибка:\n', traceback.format_exc())
        connection.commit()
        conn.commit()
        return "no"



def mod_document(document):
    current_section = document.sections[0]
    new_width, new_height = current_section.page_height, current_section.page_width
    # new_section = document.add_section(WD_SECTION.NEW_PAGE)
    # current_section = document.add_section(WD_SECTION.NEW_PAGE)
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    return


def createDocShedule(group, realGroup, students):
    columns = 26
    groupReal = realGroup
    wordDocument = docx.Document()

    style = wordDocument.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    wordDocument.add_heading(f"Журнал посещения занятий группы {groupReal}",3).alignment  = 1

    font.size = Pt(10)

    mod_document(wordDocument)
    table = wordDocument.add_table(rows=1, cols=columns)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п.п.'
    hdr_cells[1].text = 'ФИО'
    hdr_cells[1].alignment = 1
    hdr_cells[0].alignment = 1



    row = table.add_row()
    row.cells[0].merge(hdr_cells[0])
    row.cells[1].merge(hdr_cells[1])


    i = 0
    users = students
    k=0

    # table.cell(0, 2).merge(table.cell(0, 5)).text="Понедельник"
    # table.cell(0, 6).merge(table.cell(0, 9)).text="Вторник"
    # table.cell(0, 10).merge(table.cell(0, 13)).text="Среда"
    # table.cell(0, 14).merge(table.cell(0, 17)).text="Четверг"
    # table.cell(0, 18).merge(table.cell(0, 21)).text="Пятница"
    # table.cell(0, 22).merge(table.cell(0, 25)).text="Суббота"

    for i in range(columns):
        table.cell(1,i).height = Cm(5)
        table.cell(1,i).height_rule = WD_ROW_HEIGHT_RULE.AUTO
        # table.cell(1, i).text=f'\n\n\n{i}'
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "2000")
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)



    i=0
    for user in users:
        i += 1
        row = table.add_row()
        row_cells = row.cells
        # row_cells[0].height = Cm(0.3)
        # row_cells[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        row_cells[0].text = f'{i}'
        row_cells[0].width = Cm(1.19)
        row_cells[1].text = f'{user}'
        row_cells[1].width = Cm(6)


    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


    row = table.add_row()
    row_cells = row.cells
    for row in row_cells:
        row.height = Cm(5)
    # row_cells[0].merge(row_cells[1]).text = "Подпись старосты\n\n"
    row = table.add_row()
    row_cells = row.cells
    # row_cells[0].merge(row_cells[1]).text = "Подпись преподавателя\n\n"

    wordDocument.add_heading(f"Создано через бота vk.com/botraspisanie", 3).alignment = 2
    wordDocument.save("starosta_blank.docx")

def GetDocShedule(group, id, realGroup, students):
    createDocShedule(group, realGroup, students)
    a = vk.method("docs.getMessagesUploadServer", { "type" : "doc", "peer_id": id })
    b = requests.post(a["upload_url"], files= { "file" : open(str("starosta_blank")+".docx", "rb")}).json()
    c = vk.method("docs.save", {"file" : b["file"]})
    d = "doc"+str(c["doc"]["owner_id"])+"_"+str(c["doc"]["id"])
    return d
