import docx
from docx.shared import Inches, Cm
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE,WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def mod_document(document):
    current_section = document.sections[0]
    new_width, new_height = current_section.page_height, current_section.page_width
    # new_section = document.add_section(WD_SECTION.NEW_PAGE)
    # current_section = document.add_section(WD_SECTION.NEW_PAGE)
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    return


def createDocShedule(group):
    columns = 26
    groupReal = 4438
    wordDocument = docx.Document()

    style = wordDocument.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    wordDocument.add_heading(f"Журнал посещения занятий группы {groupReal}",3).alignment  = 1

    font.size = Pt(10)

    mod_document(wordDocument)
    table = wordDocument.add_table(rows=1, cols=columns)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п.п.'
    hdr_cells[1].text = 'ФИО'
    hdr_cells[1].alignment = 1
    hdr_cells[0].alignment = 1

    table.style = 'Table Grid'


    row = table.add_row()
    row.cells[0].merge(hdr_cells[0])
    row.cells[1].merge(hdr_cells[1])


    i = 0
    users = [1,2,3,4,5,6,7,8,9,1,2,3,4,5,6,7,8,9]
    k=0

    table.cell(0, 2).merge(table.cell(0, 5)).text="Понедельник"
    table.cell(0, 6).merge(table.cell(0, 9)).text="Вторник"
    table.cell(0, 10).merge(table.cell(0, 13)).text="Среда"
    table.cell(0, 14).merge(table.cell(0, 17)).text="Четверг"
    table.cell(0, 18).merge(table.cell(0, 21)).text="Пятница"
    table.cell(0, 22).merge(table.cell(0, 25)).text="Суббота"

    for i in range(columns):
        table.cell(1,i).height = Cm(5)
        table.cell(1,i).height_rule = WD_ROW_HEIGHT_RULE.AUTO
        # table.cell(1, i).text=f'\n\n\n{i}'
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "2000")
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)



    i=0
    for user in users:
        row = table.add_row()
        row_cells = row.cells
        # row_cells[0].height = Cm(0.3)
        # row_cells[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        row_cells[0].text = f'{i}'
        row_cells[0].width = Cm(1.19)
        row_cells[1].text = f'Студент{i}'
        row_cells[1].width = Cm(6)
        i+=1

    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY



    row = table.add_row()
    row_cells = row.cells
    for row in row_cells:
        row.height = Cm(5)
    row_cells[0].merge(row_cells[1]).text = "Подпись старосты\n\n"
    row = table.add_row()
    row_cells = row.cells
    row_cells[0].merge(row_cells[1]).text = "Подпись преподавателя\n\n"

    wordDocument.add_heading(f"Создано через бота vk.com/botraspisanie", 3).alignment = 2
    wordDocument.save("starosta_blank.docx")

createDocShedule(1)