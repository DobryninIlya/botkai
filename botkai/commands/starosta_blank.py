import random

import aiohttp
import requests
from bs4 import BeautifulSoup
import docx
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .. import classes as command_class
from ..classes import vk


async def mod_document(document):
    current_section = document.sections[0]
    new_width, new_height = current_section.page_height, current_section.page_width

    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    return


async def createDocShedule(group, realGroup, students):
    columns = 26
    groupReal = realGroup
    wordDocument = docx.Document()

    style = wordDocument.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    wordDocument.add_heading(f"Создано через vk.me/botraspisanie       Журнал посещения занятий группы {groupReal}",
                             3).alignment = 1

    font.size = Pt(10)

    await mod_document(wordDocument)
    table = wordDocument.add_table(rows=1, cols=columns)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п.п.'
    hdr_cells[1].text = 'ФИО'
    hdr_cells[1].alignment = 1
    hdr_cells[0].alignment = 1

    row = table.add_row()
    row.cells[0].merge(hdr_cells[0])
    row.cells[1].merge(hdr_cells[1])

    i = 0
    users = students
    k = 0

    table.cell(0, 2).merge(table.cell(0, 5)).text = "Понедельник"
    table.cell(0, 6).merge(table.cell(0, 9)).text = "Вторник"
    table.cell(0, 10).merge(table.cell(0, 13)).text = "Среда"
    table.cell(0, 14).merge(table.cell(0, 17)).text = "Четверг"
    table.cell(0, 18).merge(table.cell(0, 21)).text = "Пятница"
    table.cell(0, 22).merge(table.cell(0, 25)).text = "Суббота"

    for i in range(columns):
        table.cell(1, i).height = Cm(5)
        table.cell(1, i).height_rule = WD_ROW_HEIGHT_RULE.AUTO
        # table.cell(1, i).text=f'\n\n\n{i}'
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "2000")
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)

    i = 0
    for user in users:
        i += 1
        row = table.add_row()
        row_cells = row.cells
        # row_cells[0].height = Cm(0.3)
        # row_cells[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        row_cells[0].text = f'{i}'
        row_cells[0].width = Cm(1.19)
        row_cells[1].text = f'{user}'
        row_cells[1].width = Cm(6)

    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    row = table.add_row()
    row_cells = row.cells
    for row in row_cells:
        row.height = Cm(5)
    row_cells[0].merge(row_cells[1]).text = "Подпись старосты\n\n"
    row = table.add_row()
    row_cells = row.cells
    row_cells[0].merge(row_cells[1]).text = "Подпись преподавателя\n\n"

    wordDocument.add_heading(f"Создано через бота vk.com/botraspisanie", 3).alignment = 2
    wordDocument.save("starosta_blank.docx")


async def GetDocShedule(group, id, realGroup, students):
    await createDocShedule(group, realGroup, students)
    a = await vk.docs.getMessagesUploadServer(type="doc", peer_id=id)
    async with aiohttp.ClientSession() as session:
        async with await session.post(a["upload_url"],
                                      data={"file": open(str("starosta_blank") + ".docx", "rb")}) as response:
            b = await response.json()

    c = await vk.docs.save(file=b["file"])
    d = "doc" + str(c["doc"]["owner_id"]) + "_" + str(c["doc"]["id"])
    return d


async def info(MessageSettings, user):
    msg = "Запрос отправлен на обработку"
    msg_id = await vk.messages.send(peer_id=MessageSettings.getPeer_id(),
                                    message=msg,
                                    random_id=random.randint(1, 2147483647))
    i = 1
    async with aiohttp.ClientSession() as session:
        async with await session.post(
                "https://kai.ru/infoClick/-/info/group?id={id}".format(id=user.groupId),
                headers={'Content-Type': "application/x-www-form-urlencoded", "user-agent": "BOT RASPISANIE v.1"}) as response:
            response = await response.text()
    soup = BeautifulSoup(response, 'lxml')
    if not response:
        await vk.messages.edit(peer_id=MessageSettings.getPeer_id(),
                               message="Данные группы не найдены на сайте",
                               message_id=msg_id)
    list_students = soup.find(id="p_p_id_infoClick_WAR_infoClick10_")
    students = []
    result = ""
    for tag in list_students.find_all("td"):
        if len(tag.text) > 6:
            name_cor = (tag.text.strip().replace("\n", "").replace(
                "                                                                Староста",
                " 🟥староста")).split(" ")
            name = ""
            try:
                name = name_cor[0] + " " + name_cor[1][0].capitalize() + "." + name_cor[2][0].capitalize() + "."
            except:
                name = name_cor[0][:20] + ". "
            try:
                name += name_cor[3]
            except:
                pass
            students.append(name)

    try:
        att = await GetDocShedule(user.groupId, MessageSettings.getPeer_id(), user.RealGroup, students)
    except:
        pass
    await vk.messages.send(peer_id=MessageSettings.getPeer_id(),
                           message="Бланк посещения",
                           attachment=att, random_id=random.randint(1, 2147483647))
                           # message_id=msg_id)

    return "ok"


command = command_class.Command()

command.keys = ['бланк посещения', "журнал посещения"]
command.desciption = 'отображение полного списка группы'
command.process = info
command.payload = "starosta_blank"
