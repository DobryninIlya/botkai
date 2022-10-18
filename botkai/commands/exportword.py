import aiohttp

from .. import classes as command_class
from ..classes import vk, cursor, connection
import random
import json
from ..keyboards import submenu

import docx
import requests
import datetime
import traceback

today = datetime.date.today()
BASE_URL = 'https://kai.ru/raspisanie'

class ShedRow(object):
    def __init__(self, dayTime, dayDate, disciplName, disciplType, audNum, buildNum, prepodName):
        self.dayTime=dayTime
        self.dayDate = dayDate
        self.disciplName = disciplName
        self.disciplType = disciplType
        self.audNum = audNum
        self.buildNum = buildNum
        self.prepodName = prepodName


async def getResponse(groupId):
    sql = "SELECT * FROM saved_timetable WHERE groupp = {}".format(groupId)
    cursor.execute(sql)
    result = cursor.fetchone()
    if result == None:
        try:
            async with aiohttp.ClientSession() as session:
                async with await session.post(BASE_URL, data="groupId=" + str(groupId),
                                              headers={'Content-Type': "application/x-www-form-urlencoded"},
                                              params={"p_p_id": "pubStudentSchedule_WAR_publicStudentSchedule10",
                                                      "p_p_lifecycle": "2", "p_p_resource_id": "schedule"},
                                              timeout=3) as response:
                    response = await response.json()
            sql = "INSERT INTO saved_timetable VALUES ({}, '{}', '{}')".format(groupId, datetime.date.today(),
                                                                               json.dumps(response))
            cursor.execute(sql)
            connection.commit()
            return True, response
        except ConnectionError as err:
            return False, "&#9888;Ошибка подключения к серверу типа ConnectionError. Вероятно, сервера КАИ были выведены из строя.&#9888;"
        except aiohttp.ServerTimeoutError as err:
            return False, "&#9888;Ошибка подключения к серверу типа Timeout. Вероятно, сервера КАИ перегружены.&#9888;"
        except:
            return False, ""

    else:
        date_update = result[1]
        timetable = result[2]
        if date_update + datetime.timedelta(days=2) < today:
            try:
                async with aiohttp.ClientSession() as session:
                    async with await session.post(BASE_URL, data="groupId=" + str(groupId),
                                                  headers={'Content-Type': "application/x-www-form-urlencoded"},
                                                  params={"p_p_id": "pubStudentSchedule_WAR_publicStudentSchedule10",
                                                          "p_p_lifecycle": "2", "p_p_resource_id": "schedule"},
                                                  timeout=3) as response:
                        response = await response.json()
                sql = "UPDATE saved_timetable SET shedule = '{}', date_update = '{}' WHERE groupp = {}".format(
                    json.dumps(response), datetime.date.today(), groupId)
                cursor.execute(sql)
                connection.commit()
                return True, response
            except:
                sql = "SELECT shedule FROM saved_timetable WHERE groupp = {}".format(groupId)
                cursor.execute(sql)
                result = cursor.fetchone()[0]
                return True, json.loads(result)
        else:
            sql = "SELECT shedule FROM saved_timetable WHERE groupp = {}".format(groupId)
            cursor.execute(sql)
            result = cursor.fetchone()[0]
            return True, json.loads(result)
    return




    return 

async def TimetableWrite(groupId):
    isNormal, response = await getResponse(groupId)
    if not isNormal:
        return response

    rows = 0
    lis = []

    elem = response[str(1)]
    try:
        lis.append("Понедельник")
        for day in elem:
            lis.append(ShedRow(day["dayTime"], day["dayDate"], day["disciplName"], day["disciplType"], day["audNum"], day["buildNum"], day["prepodName"]))
        elem = response[str(2)]
        lis.append("Вторник")
        for day in elem:
            lis.append(ShedRow(day["dayTime"], day["dayDate"], day["disciplName"], day["disciplType"], day["audNum"], day["buildNum"], day["prepodName"]))
        rows = len(lis)
        elem = response[str(3)]
        lis.append("Среда")
        for day in elem:
            lis.append(ShedRow(day["dayTime"], day["dayDate"], day["disciplName"], day["disciplType"], day["audNum"], day["buildNum"], day["prepodName"]))
        rows = len(lis)
        elem = response[str(4)]
        lis.append("Четверг")
        for day in elem:
            lis.append(ShedRow(day["dayTime"], day["dayDate"], day["disciplName"], day["disciplType"], day["audNum"], day["buildNum"], day["prepodName"]))
        rows = len(lis)
        elem = response[str(5)]
        lis.append("Пятница")
        for day in elem:
            lis.append(ShedRow(day["dayTime"], day["dayDate"], day["disciplName"], day["disciplType"], day["audNum"], day["buildNum"], day["prepodName"]))
        rows = len(lis)
        elem = response[str(6)]
        lis.append("Суббота")
        for day in elem:
            lis.append(ShedRow(day["dayTime"], day["dayDate"], day["disciplName"], day["disciplType"], day["audNum"], day["buildNum"], day["prepodName"]))
    except:
        pass
    return lis

async def createDocShedule(group):

    wordDocument = docx.Document("blank.docx")


    lis = await TimetableWrite(group)
    for day in lis:
        if day in ["Понедельник","Вторник","Среда","Четверг","Пятница","Суббота"]:
            par = wordDocument.add_heading(day, 3)
            par.bold = True

        
        else:
            par = wordDocument.add_paragraph((str(day.dayTime)).rstrip() + " " + ((str(day.dayDate)).rstrip()).ljust(8) + " " + str(day.disciplName) + " " + (str(day.disciplType)).upper() + " " + (str(day.audNum)).rstrip() + " ауд  " + (str(day.buildNum)).rstrip() + "зд.  " + (str(day.prepodName)).rstrip())
            # par.style = "No Spacing"
    wordDocument.save(str(group)+".docx")

async def GetDocShedule(group, id):
    await createDocShedule(group)
    a = await vk.docs.getMessagesUploadServer(type="doc", peer_id=id)
    async with aiohttp.ClientSession() as session:
        async with await session.post(a["upload_url"],
                                      data={"file": open(str(group)+".docx", "rb")}) as response:
            b = await response.json()

    c = await vk.docs.save(file=b["file"])
    d = "doc" + str(c["doc"]["owner_id"]) + "_" + str(c["doc"]["id"])
    return d


async def info(MessageSettings, user):
    await vk.messages.send(peer_id=MessageSettings.getPeer_id(),
                           message="Твое персональное расписание",
                           random_id=random.randint(1, 2147483647),
                           keyboard=submenu,
                           attachment=await GetDocShedule(UserParams.groupId, MessageSettings.getPeer_id()))


info_command = command_class.Command()

info_command.keys = ['экспорт в word']
info_command.desciption = 'доп меню'
info_command.payload = "exportword"
info_command.process = info
